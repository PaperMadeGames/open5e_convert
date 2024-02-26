import pandas as pd
import requests
import sys
import argparse
from docx import Document

def speed_to_text(speed):
    for k, v in speed.items():
        if k not in ['walk', 'burrow', 'climb', 'fly', 'swim', 'hover', 'bur.', 'lightwalking', 'notes']:
            sys.exit(f'invalid speed key found: {k}')
    speed_text = ""
    if 'walk' in speed:
        speed_text = speed_text + f"{speed['walk']} ft."
    if 'burrow' in speed:
        if speed_text != "":
            speed_text = speed_text + ", "
        speed_text = speed_text + f"Burrow {speed['burrow']} ft."
    if 'climb' in speed:
        if speed_text != "":
            speed_text = speed_text + ", "
        speed_text = speed_text + f"Climb {speed['climb']} ft."
    if 'fly' in speed:
        if speed_text != "":
            speed_text = speed_text + ", "
        speed_text = speed_text + f"Fly {speed['fly']} ft.{' (hover)' if 'hover' in speed and speed['hover'] else ''}"
    if 'swim' in speed:
        if speed_text != "":
            speed_text = speed_text + ", "
        speed_text = speed_text + f"Swim {speed['swim']} ft."
    if 'bur.' in speed:
        if speed_text != "":
            speed_text = speed_text + ", "
        speed_text = speed_text + f"Burrow {speed['bur.']} ft."
    if 'lightwalking' in speed:
        if speed_text != "":
            speed_text = speed_text + ", "
        speed_text = speed_text + f"Lightwalking {speed['lightwalking']} ft."
    if 'notes' in speed:
        if speed_text != "":
            speed_text = speed_text + ", "
        speed_text = speed_text + f"Notes {speed['notes']} ft."
    return speed_text

def attr_to_text(attr):
    modifier = int(attr/2) - 5
    return f"{attr} ({modifier:+})"

def save_to_text(monster):
    save_text = ""
    if monster['strength_save']:
        save_text = save_text + f"Str {monster['strength_save']:+}"
    if monster['dexterity_save']:
        if save_text != "":
            save_text = save_text + ", "
        save_text = save_text + f"Dex {monster['dexterity_save']:+}"
    if monster['constitution_save']:
        if save_text != "":
            save_text = save_text + ", "
        save_text = save_text + f"Con {monster['constitution_save']:+}"
    if monster['intelligence_save']:
        if save_text != "":
            save_text = save_text + ", "
        save_text = save_text + f"Int {monster['intelligence_save']:+}"
    if monster['wisdom_save']:
        if save_text != "":
            save_text = save_text + ", "
        save_text = save_text + f"Wis {monster['wisdom_save']:+}"
    if monster['charisma_save']:
        if save_text != "":
            save_text = save_text + ", "
        save_text = save_text + f"Cha {monster['charisma_save']:+}"
    return save_text

def setup_document(title=""):
    # Create a new Word document
    doc = Document()
    doc.add_heading(title, level=0)
    return(doc)

def monster_import(results, doc):
    doc.add_heading("Monsters", level=1)
    for monster in results:
        print('Importing ' + str(monster['name']))
        doc.add_heading(str(monster['name']), level=2)
        doc.add_paragraph(f"{monster['size'] if monster['size'].lower() != 'titanic' else 'Gargantuan'} {monster['type']}, {monster['alignment']}")
        doc.add_paragraph(f"Armor Class {monster['armor_class']} {monster['armor_desc'] if (monster['armor_desc']) else ''}")
        doc.add_paragraph(f"Hit Points {monster['hit_points']} ({monster['hit_dice']})")
        doc.add_paragraph("Speed " + speed_to_text(monster['speed']))
        doc.add_paragraph("STR DEX CON INT WIS CHA")
        doc.add_paragraph(attr_to_text(monster['strength']) + " " +
                attr_to_text(monster['dexterity']) + " " +
                attr_to_text(monster['constitution']) + " " +
                attr_to_text(monster['intelligence']) + " " +
                attr_to_text(monster['wisdom']) + " " +
                attr_to_text(monster['charisma']) )
        save_text = save_to_text(monster)
        if save_text != "":
            doc.add_paragraph("Saving Throws " + save_to_text(monster))
        # doc.add_paragraph("Skills") - they don't seem to be used
        doc.add_paragraph("Damage Vulnerabilities  " + str(monster['damage_vulnerabilities']))
        doc.add_paragraph("Damage Resistances  " + str(monster['damage_resistances']))
        doc.add_paragraph("Damage Immunities " + str(monster['damage_immunities']))
        doc.add_paragraph("Condition Immunities " + str(monster['condition_immunities']))
        doc.add_paragraph("Senses " + str(monster['senses']))
        doc.add_paragraph("Languages " + str(monster['languages']))
        doc.add_paragraph("Challenge " + str(monster['challenge_rating']))

        if monster['special_abilities']:
            for index, a in enumerate(monster['special_abilities']):
                p = doc.add_paragraph("")
                runner = p.add_run(a['name'] + ". ")
                runner.bold = True
                runner.italic = True
                p.add_run(a['desc'])
                
        if monster['actions']:
            doc.add_heading("Actions", level=3)
            for index, a in enumerate(monster['actions']):
                p = doc.add_paragraph("")
                runner = p.add_run(a['name'] + ". ")
                runner.bold = True
                runner.italic = True
                p.add_run(a['desc'])

        if monster['bonus_actions']:
            doc.add_heading("Bonus Actions", level=3)
            for index, a in enumerate(monster['bonus_actions']):
                p = doc.add_paragraph("")
                runner = p.add_run(a['name'] + ". ")
                runner.bold = True
                runner.italic = True
                p.add_run(a['desc'])

        if monster['reactions']:
            doc.add_heading("Reactions", level=3)
            for index, a in enumerate(monster['reactions']):
                p = doc.add_paragraph("")
                runner = p.add_run(a['name'] + ". ")
                runner.bold = True
                runner.italic = True
                p.add_run(a['desc'])

        if monster['legendary_actions']:
            doc.add_heading("Legendary Actions", level=3)
            if monster['legendary_desc']:
                doc.add_paragraph(monster['legendary_desc'])
            for index, a in enumerate(monster['legendary_actions']):
                p = doc.add_paragraph("")
                runner = p.add_run(a['name'] + ". ")
                runner.bold = True
                runner.italic = True
                p.add_run(a['desc'])

def spell_import(results, doc):
    doc.add_heading("Spells", level=1)
    for spell in results:
        print('Importing ' + str(spell['name']))
        doc.add_heading(str(spell['name']), level=2)
        if (spell['level'] == 'Cantrip'):
            doc.add_paragraph(f"{spell['school']} {spell['level']}")
        else:
            doc.add_paragraph(f"{spell['level']} {spell['school']} {'(ritual}' if spell['ritual'] == 'yes' else ''}")
        doc.add_paragraph(f"Casting Time {spell['casting_time']}")
        doc.add_paragraph(f"Range {spell['range']}")
        doc.add_paragraph(f"Components {spell['components']}{' (' + spell['material'] + ')' if spell['material'] != '' else ''}")
        doc.add_paragraph(f"Duration {'concentration, ' if spell['concentration'] == 'yes' else '' }{spell['duration']}")
        doc.add_paragraph(f"{'classes ' + spell['dnd_class'] if spell['dnd_class'] != '' else ''}")
        doc.add_paragraph(spell['desc'])
        doc.add_paragraph(spell['higher_level'])

    doc.add_heading("Artificer", level=1)
    for spell in results:
        if (spell['dnd_class'].find('Artificer') >= 0):
            doc.add_paragraph(spell['name'])
    doc.add_heading("Wizard", level=1)
    for spell in results:
        if (spell['dnd_class'].find('Wizard') >= 0):
            doc.add_paragraph(spell['name'])
    doc.add_heading("Sorcerer", level=1)
    for spell in results:
        if (spell['dnd_class'].find('Sorceror') >= 0 or spell['dnd_class'].find('Sorcerer') >= 0):
            doc.add_paragraph(spell['name'])
    doc.add_heading("Cleric", level=1)
    for spell in results:
        if (spell['dnd_class'].find('Cleric') >= 0):
            doc.add_paragraph(spell['name'])
    doc.add_heading("Ranger", level=1)
    for spell in results:
        if (spell['dnd_class'].find('Ranger') >= 0):
            doc.add_paragraph(spell['name'])
    doc.add_heading("Warlock", level=1)
    for spell in results:
        if (spell['dnd_class'].find('Warlock') >= 0):
            doc.add_paragraph(spell['name'])
    doc.add_heading("Bard", level=1)
    for spell in results:
        if (spell['dnd_class'].find('Bard') >= 0):
            doc.add_paragraph(spell['name'])
    doc.add_heading("Paladin", level=1)
    for spell in results:
        if (spell['dnd_class'].find('Paladin') >= 0 or spell['dnd_class'].find('Herald') >= 0):
            doc.add_paragraph(spell['name'])
    doc.add_heading("Druid", level=1)
    for spell in results:
        if (spell['dnd_class'].find('Druid') >= 0):
            doc.add_paragraph(spell['name'])


def item_import(results, doc):
    doc.add_heading("Items", level=1)
    for item in results:
        print('Importing ' + str(item['name']))
        doc.add_heading(str(item['name']), level=2)
        doc.add_paragraph(f"{item['type']}, {item['rarity']} {'(' + item['requires_attunement'] + ')' if item['requires_attunement'] != '' else ''}")
        doc.add_paragraph(item['desc'])

def feat_import(results, doc):
    doc.add_heading("Feats", level=1)
    for feat in results:
        print('Importing ' + str(feat['name']))
        doc.add_heading(str(feat['name']), level=2)
        if (feat['prerequisite'] != ''):
            doc.add_paragraph(f"Prerequisite {feat['prerequisite']}")
        doc.add_paragraph(feat['desc'])

def main():
    argv = sys.argv[1:]

    parser = argparse.ArgumentParser(
                        description='Convert open5e content into Docx format',
                        epilog='You can do it!!')
    parser.add_argument('-f', '--feature', required=True,
                        help='the feature type you are interested in fetching from open5e',
                        choices=['monsters', 'spells', 'magicitems', 'feats'])
    parser.add_argument('-s', '--source', required=True,
                        help='the source of the data you want to import from open5e',
                        choices=['menagerie', 'wotc-srd', 'a5e', 'dmag', 'dmag-e', 'warlock', 'kp', 'toh', 'tob', 'tob2', 'tob3', 'cc', 'taldorei', 'vom'])
    parser.add_argument('-t', '--title', dest='doc_title', default='', help='the title of the document you want to create')
    parser.add_argument('-o', '--out', dest='doc_filename', default='', help='the name of the output file')
    args = parser.parse_args()

    source_name = ""
    if (args.source == 'menagerie'):
        source_name = 'Level Up Advanced 5e Monstrous Menagerie'
    elif (args.source == 'wotc-srd'):
        source_name = '5e Core Rules'
    elif (args.source == 'a5e'):
        source_name = 'Level Up Advanced 5e'
    elif (args.source == 'dmag'):
        source_name = 'Deep Magic 5e'
    elif (args.source == 'dmag-e'):
        source_name = 'Deep Magic Extended'
    elif (args.source == 'warlock'):
        source_name = 'Warlock Archives'
    elif (args.source == 'kp'):
        source_name = 'Kobold Press Compilation'
    elif (args.source == 'toh'):
        source_name = 'Tome of Heroes'
    elif (args.source == 'tob'):
        source_name = 'Tome of Beasts'
    elif (args.source == 'tob2'):
        source_name = 'Tome of Beasts 2'
    elif (args.source == 'tob3'):
        source_name = 'Tome of Beasts 3'
    elif (args.source == 'cc'):
        source_name = 'Creature Codex'
    elif (args.source == 'taldorei'):
        source_name = 'Critical Role: Tal’Dorei Campaign Setting'
    elif (args.source == 'vom'):
        source_name = 'Vault of Magic'
    else:
        sys.exit(f"Invalid source {args.source}. Valid sources are " +
                    "menagerie, wotc-srd, a5e, dmag, dmag-e, warlock, kp, toh, tob, tob2, tob3, cc, taldorei, vom.")
    if (args.doc_title == ""):
        args.doc_title = f"{args.source.capitalize() if source_name == '' else source_name} {args.feature.capitalize()}"
        print(f"No title provided, using default title '{args.doc_title}'")
    if (args.doc_filename == ''):
        args.doc_filename = f"{args.doc_title}.docx"
        print(f"No filename provided, using default filename '{args.doc_filename}'")

    doc = setup_document(title=args.doc_title)

    # load data from API
    page_no = 0
    params = {'document__slug__iexact': args.source, 'page': page_no}
    results = []
    sys.stdout.write(f"Downloading https://api.open5e.com/v1/{args.feature}")
    while True:
        params['page'] += 1
        try:
            sys.stdout.write(".")
            sys.stdout.flush()
            r = requests.get(f"https://api.open5e.com/v1/{args.feature}", params=params)
            r.raise_for_status()
        except requests.exceptions.HTTPError as err:
            break

        df = pd.DataFrame(r.json())
        for x in df['results'].items():
            results.append(x[1])
    sys.stdout.write("Done!\n")

    if (args.feature == 'monsters'):
        print("Importing Monsters")
        monster_import(results, doc)
    elif (args.feature == 'spells'):
        print("Importing Spells")
        spell_import(results, doc)
    elif (args.feature == 'magicitems'):
        print("Importing Magic Items")
        item_import(results, doc)
    elif (args.feature == 'feats'):
        print("Importing feats")
        feat_import(results, doc)


    # Save the Word document
    doc.save(args.doc_filename)
    print ("wrote file " + args.doc_filename)


main()
